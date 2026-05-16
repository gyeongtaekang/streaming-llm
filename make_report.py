from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 페이지 여백
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

def set_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def heading(doc, text, level=1, size=16, color=(0,70,127)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    return p

def body(doc, text, size=11, indent=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def shade_row(row, hex_color="D9E2F3"):
    for cell in row.cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'),   'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'),  hex_color)
        cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = t.rows[0]
    shade_row(hrow, "2E74B5")
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255,255,255))
    for ri, row_data in enumerate(rows):
        drow = t.rows[ri + 1]
        if ri % 2 == 1:
            shade_row(drow, "EBF3FB")
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size=10)
    if col_widths:
        for ri2 in range(len(t.rows)):
            for ci2, w in enumerate(col_widths):
                t.rows[ri2].cells[ci2].width = Cm(w)
    return t

# ══════════════════════════════════════════
# 표지
# ══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("ROCStories Fine-tuning +\nSelfCheckGPT NLI 평가 결과 보고서")
set_font(run, size=20, bold=True, color=(0,70,127))

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run(f"작성일: {datetime.date.today().strftime('%Y년 %m월 %d일')}")
set_font(run, size=11, color=(89,89,89))

doc.add_page_break()

# ══════════════════════════════════════════
# 1. 실험 개요
# ══════════════════════════════════════════
heading(doc, "1. 실험 개요")
body(doc, "본 실험은 GPT-2 XL 모델을 ROCStories 데이터셋으로 파인튜닝한 뒤, "
          "SelfCheckGPT(NLI 방식)를 이용해 Story Cloze Task의 정확도를 측정하였습니다.")

heading(doc, "1.1 연구 목적", level=2, size=13, color=(31,73,125))
items = [
    "GPT-2 XL(1.5B 파라미터)을 5문장 단편 소설 데이터로 파인튜닝",
    "SelfCheckGPT NLI 방식으로 모델의 사실 일관성(factual consistency) 측정",
    "Story Cloze Task(올바른 결말 선택)에서 랜덤 기준선(50%) 대비 성능 향상 확인",
]
for it in items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(it)
    set_font(run, size=11)

# ══════════════════════════════════════════
# 2. 실험 환경
# ══════════════════════════════════════════
heading(doc, "2. 실험 환경")

heading(doc, "2.1 하드웨어", level=2, size=13, color=(31,73,125))
add_table(doc,
    ["항목", "사양"],
    [
        ["GPU", "NVIDIA RTX A5000 × 3"],
        ["GPU VRAM", "24.6 GB × 3"],
        ["학습 중 VRAM 사용량", "~18.1 GB / GPU"],
        ["학습 방식", "DDP (Distributed Data Parallel)"],
    ],
    col_widths=[5, 10]
)

doc.add_paragraph()
heading(doc, "2.2 소프트웨어", level=2, size=13, color=(31,73,125))
add_table(doc,
    ["라이브러리", "버전 / 설정"],
    [
        ["PyTorch",        "2.x + CUDA"],
        ["Transformers",   "5.8.0"],
        ["Accelerate",     "1.13.0 (mixed_precision=bf16)"],
        ["bitsandbytes",   "0.49.2 (8-bit AdamW)"],
        ["selfcheckgpt",   "NLI variant (DeBERTa-v3-large)"],
        ["sentencepiece",  "0.2.1"],
    ],
    col_widths=[5, 10]
)

# ══════════════════════════════════════════
# 3. 데이터셋
# ══════════════════════════════════════════
doc.add_paragraph()
heading(doc, "3. 데이터셋")
add_table(doc,
    ["분류", "파일", "샘플 수"],
    [
        ["학습 (Train)", "ROCStories spring2016 + winter2017", "98,000"],
        ["검증 (Val)",   "Story Cloze val spring2016 + winter2018", "300 (샘플링)"],
    ],
    col_widths=[3, 10, 3]
)
doc.add_paragraph()
body(doc, "※ ROCStories: 5문장으로 구성된 일상적 단편 소설 데이터셋 (문장1~5).", size=10, indent=0.5)
body(doc, "※ Story Cloze Task: 문장1~4(맥락) + 두 개의 후보 결말 중 올바른 것을 선택(AnswerRightEnding=1 or 2).", size=10, indent=0.5)

# ══════════════════════════════════════════
# 4. 학습 설정
# ══════════════════════════════════════════
heading(doc, "4. 학습 설정 (Phase 1: Fine-tuning)")
add_table(doc,
    ["하이퍼파라미터", "값", "설명"],
    [
        ["모델",           "GPT-2 XL",    "1,558M 파라미터, BF16"],
        ["GPU 수",         "3",           "DDP, accelerate launch"],
        ["GPU당 배치",     "16",          ""],
        ["Gradient Accum","4",           "Effective batch = 16×3×4 = 192"],
        ["최대 시퀀스",    "512 tokens",  ""],
        ["학습률",         "3e-5",        "Linear warmup (10%)"],
        ["에폭",           "1",           ""],
        ["옵티마이저",     "8-bit AdamW", "bitsandbytes"],
        ["Grad Checkpt",  "ON",          "use_reentrant=False"],
        ["Precision",     "BF16",        "mixed precision"],
    ],
    col_widths=[4, 4, 7]
)

doc.add_paragraph()
body(doc, "학습 소요 시간: 약 36분 (3-GPU DDP 기준; 단일 GPU 대비 ~3배 빠름)")

# ══════════════════════════════════════════
# 5. SelfCheckGPT 방법론
# ══════════════════════════════════════════
heading(doc, "5. SelfCheckGPT NLI 평가 방법론 (Phase 2)")
body(doc, "SelfCheckGPT(NLI)는 LLM이 생성한 여러 샘플 응답 간의 일관성을 NLI(자연어 추론) 모델로 측정합니다.")
doc.add_paragraph()

steps = [
    ("Step 1 — 샘플 생성",
     "파인튜닝된 GPT-2 XL로 각 맥락(문장1~4)에 대해 5개의 응답을 샘플링\n"
     "(greedy 1개 + temperature=1.0, top_p=0.95 샘플링 5개, max_new_tokens=80)"),
    ("Step 2 — NLI 점수 계산",
     "DeBERTa-v3-large NLI 모델로 각 quiz 문장과 5개 샘플 간의\n"
     "contradiction 확률(P(contradiction))을 계산"),
    ("Step 3 — 예측",
     "NLI contradiction 점수가 낮을수록 → 샘플들과 더 일관성 있는 문장\n"
     "→ score(Q1) ≤ score(Q2)이면 Q1 선택, 그렇지 않으면 Q2 선택"),
]
for title, desc in steps:
    heading(doc, title, level=2, size=12, color=(31,73,125))
    body(doc, desc, size=11, indent=0.5)

# ══════════════════════════════════════════
# 6. 결과
# ══════════════════════════════════════════
heading(doc, "6. 평가 결과")

heading(doc, "6.1 최종 정확도", level=2, size=13, color=(31,73,125))
add_table(doc,
    ["평가 지표", "값"],
    [
        ["Story Cloze Accuracy", "0.6533  (196 / 300)"],
        ["Random Baseline",      "0.5000  (150 / 300)"],
        ["향상폭",               "+0.1533  (+15.3%p)"],
        ["평가 샘플 수",         "300"],
        ["SelfCheck 샘플 수",    "5 per example"],
    ],
    col_widths=[6, 9]
)

doc.add_paragraph()
heading(doc, "6.2 구간별 정확도 추이", level=2, size=13, color=(31,73,125))
add_table(doc,
    ["평가 구간", "누적 정확도"],
    [
        ["50  / 300", "0.6400"],
        ["100 / 300", "0.6500"],
        ["150 / 300", "0.6933"],
        ["200 / 300", "0.6750"],
        ["250 / 300", "0.6640"],
        ["300 / 300", "0.6533"],
    ],
    col_widths=[5, 5]
)

doc.add_paragraph()
heading(doc, "6.3 예측 예시", level=2, size=13, color=(31,73,125))

examples = [
    {
        "context": "Pamela was so excited to finally have a boyfriend, Ralph. Ralph never listened to what she said. He rarely showed up when he said he would...",
        "q1": "Ralph then asked Pamela to marry her.",
        "q2": "Pamela decided to break up with Ralph.",
        "q1_score": 0.9951, "q2_score": 0.9713,
        "pred": 2, "ans": 2, "correct": True,
    },
    {
        "context": "Reg was hoping for a snow day. There was a bad storm overnight. He got his wish and school was canceled! He spent the whole day playing...",
        "q1": "Reg hoped it would never snow again.",
        "q2": "Reg hoped he would get another snow day soon.",
        "q1_score": 0.9798, "q2_score": 0.9385,
        "pred": 2, "ans": 2, "correct": True,
    },
    {
        "context": "Jackie had been a long distance runner for many Year's. Despite her very full iPod, she became bored with music...",
        "q1": "She listened to a book while she ran.",
        "q2": "She was glad she is never bored of music.",
        "q1_score": 0.9687, "q2_score": 0.9241,
        "pred": 2, "ans": 1, "correct": False,
    },
]

for ex in examples:
    mark = "✓ 정답" if ex["correct"] else "✗ 오답"
    color = (0,112,0) if ex["correct"] else (192,0,0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"[{mark}]")
    set_font(run, size=11, bold=True, color=color)

    body(doc, f"맥락: {ex['context']}", size=10, indent=0.5)
    body(doc, f"Quiz 1 (NLI={ex['q1_score']:.4f}): {ex['q1']}", size=10, indent=0.5)
    body(doc, f"Quiz 2 (NLI={ex['q2_score']:.4f}): {ex['q2']}", size=10, indent=0.5)
    body(doc, f"예측={ex['pred']}  정답={ex['ans']}", size=10, indent=0.5, space_after=2)

# ══════════════════════════════════════════
# 7. 결론
# ══════════════════════════════════════════
heading(doc, "7. 결론 및 고찰")
conclusions = [
    "GPT-2 XL을 ROCStories로 파인튜닝한 결과, Story Cloze Task에서 "
    "SelfCheckGPT NLI 기준 65.3%의 정확도를 달성하였습니다 (랜덤 기준선 50% 대비 +15.3%p).",
    "3-GPU DDP(Distributed Data Parallel) + Gradient Checkpointing + 8-bit AdamW를 조합하여 "
    "GPU당 18.1GB VRAM으로 안정적인 학습을 완료하였습니다.",
    "SelfCheckGPT는 별도의 레이블 없이 모델 자체의 일관성만으로 평가하는 비지도 방식임에도 "
    "유의미한 성능 향상을 보였습니다.",
    "에폭 수 증가, 프롬프트 설계 개선, 더 큰 모델 사용 등을 통해 추가적인 성능 향상 가능성이 있습니다.",
]
for c in conclusions:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(c)
    set_font(run, size=11)

# ══════════════════════════════════════════
# 저장
# ══════════════════════════════════════════
out_path = "ROCStories_SelfCheckGPT_결과보고서.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
